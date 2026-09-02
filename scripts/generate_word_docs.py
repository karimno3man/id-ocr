"""Generate project Word documentation (technical + user guide)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"


def set_normal_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_title_page(document: Document, title: str, subtitle: str) -> None:
    document.add_paragraph()
    heading = document.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = document.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org = document.add_paragraph("iSchool")
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    document.add_paragraph()


def build_technical_doc() -> Document:
    doc = Document()
    set_normal_style(doc)
    add_title_page(
        doc,
        "Egyptian National ID OCR",
        "Technical Documentation",
    )

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This project extracts structured data from Egyptian national ID cards using a hybrid "
        "computer-vision pipeline. The system detects 16 labeled field regions on the card, reads "
        "Arabic text and Eastern Arabic digits, decodes the 14-digit national ID number, and presents "
        "editable results through a web interface."
    )
    doc.add_paragraph(
        "The solution combines YOLO object detection (Ultralytics), a specialized digit detector "
        "for the ID number field, PaddleOCR for Arabic text recognition, and a rule-based Egyptian "
        "NID decoder."
    )

    doc.add_heading("2. Problem Definition", level=1)
    doc.add_paragraph(
        "Egyptian national ID cards contain Arabic script, Eastern Arabic numerals (٠–٩), and a "
        "fixed multi-region layout split across the front and back of the card. Manual transcription "
        "is slow and error-prone. The goal is to automate field localization and text extraction while "
        "keeping a human review step for corrections."
    )
    bullets = [
        "Variable photo quality (blur, glare, rotation, partial crops)",
        "Mixed Arabic text and numeric fields",
        "ID number requires per-digit detection, not general OCR alone",
        "Some fields appear only on the back of the card",
    ]
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. System Architecture", level=1)
    doc.add_paragraph("The end-to-end pipeline has four stages:")
    stages = [
        ("Stage 1 — Field localization", "A YOLO model detects bounding boxes for all 16 field classes on front and back images."),
        ("Stage 2 — ID digit reading", "The ID crop is passed to a 10-class digit YOLO model. Detected digits are sorted left-to-right into a 14-digit string."),
        ("Stage 3 — Text OCR", "The remaining 15 fields are read with PaddleOCR (Arabic), including RTL line stitching and field-specific preprocessing."),
        ("Stage 4 — Post-processing", "Front and back results are merged. The 14-digit ID is decoded into birth date, governorate, and gender metadata."),
    ]
    for title, body in stages:
        doc.add_paragraph(title, style="List Number")
        doc.add_paragraph(body)

    doc.add_paragraph("Main code modules:")
    add_table(
        doc,
        ["Module", "Path", "Role"],
        [
            ["CLI pipeline", "yolo+ocr/run_pipeline.py", "Batch inference on image folders"],
            ["Card extractor", "yolo+ocr/card_extractor.py", "Shared inference for CLI and web API"],
            ["Field OCR", "yolo+ocr/field_ocr.py", "PaddleOCR wrapper for Arabic text fields"],
            ["Digit NID", "yolo+ocr/digit_nid.py", "Per-digit ID reading"],
            ["NID decoder", "yolo+ocr/decode_nid.py", "Parse 14-digit Egyptian ID structure"],
            ["Web API", "web/app.py", "FastAPI endpoints for extract and submit"],
            ["Web UI", "web/static/", "Upload, review, and edit interface"],
        ],
    )

    doc.add_heading("4. Technologies Used", level=1)
    add_table(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Deep learning", "PyTorch", "Training and inference"],
            ["Object detection", "Ultralytics YOLO (YOLO26m, YOLOv8n)", "Field and digit detection"],
            ["OCR", "PaddleOCR 3.x / PaddlePaddle 3.2.2", "Arabic text recognition"],
            ["Vision", "OpenCV, NumPy, Pillow", "Image decode, crops, preprocessing"],
            ["Web", "FastAPI, Uvicorn", "REST API and static UI serving"],
            ["Tracking", "MLflow (SQLite)", "Experiment logging during training"],
            ["Data", "Roboflow Universe", "Labeled datasets (YOLOv8 export)"],
            ["Optional export", "Google Apps Script", "Append reviewed rows to Google Sheets"],
        ],
    )

    doc.add_heading("5. Datasets", level=1)
    doc.add_heading("5.1 Field Localization — Thndr National Card", level=2)
    doc.add_paragraph(
        "Source: Roboflow Universe — Thndr National Card (Thndr-National-Card.v4-v4.yolov8). "
        "Format: YOLOv8 object detection. Not committed to git; downloaded locally for training."
    )
    doc.add_paragraph("16 classes:")
    doc.add_paragraph(
        "Add1, Add2, Back, ExpDate, First_Name, Front, Gender, HusbandName, ID, IssueDate, "
        "Job1, Job2, Last_Name, Religion, Serial_Num, Status"
    )
    add_table(
        doc,
        ["Split", "Images", "Notes"],
        [
            ["Train", "4,071", "42 corrupt labels excluded during training"],
            ["Valid", "194", "1 corrupt label excluded"],
            ["Test", "Present", "Defined in dataset YAML"],
        ],
    )

    doc.add_heading("5.2 Digit Detection — cro4", level=2)
    doc.add_paragraph(
        "Source: Roboflow Universe — cro4 (cro4.v1-8.yolov8). "
        "One bounding box per Eastern Arabic digit (classes 0–9). Split: 80% train / 10% valid / 10% test."
    )
    add_table(
        doc,
        ["Split", "Images", "Digit boxes"],
        [
            ["Train", "6,795", "~95,130"],
            ["Valid", "849", "~11,886"],
            ["Test", "Present", "Per dataset split"],
        ],
    )

    doc.add_heading("6. Model Training", level=1)
    doc.add_heading("6.1 Localization Model (train_nid_yolo.ipynb)", level=2)
    add_table(
        doc,
        ["Parameter", "Value"],
        [
            ["Base model", "YOLO26m (yolo26m.pt)"],
            ["Classes", "16 field regions"],
            ["Image size", "640 px"],
            ["Epochs", "50"],
            ["Batch size", "8"],
            ["Augmentation", "Horizontal flip enabled (fliplr=0.5)"],
            ["MLflow experiment", "nid-localization"],
            ["Output weights", "runs/nid_localize_<timestamp>/weights/best.pt"],
        ],
    )

    doc.add_heading("6.2 Digit Model (train_nid_digits.ipynb)", level=2)
    add_table(
        doc,
        ["Parameter", "Value"],
        [
            ["Base model", "YOLOv8n (yolov8n.pt)"],
            ["Classes", "10 digits (0–9)"],
            ["Image size", "640 px"],
            ["Epochs", "50"],
            ["Batch size", "32"],
            ["Augmentation", "fliplr=0 (digits are left-to-right)"],
            ["MLflow experiment", "nid-digits"],
            ["Output weights", "runs/nid_digits/weights/best.pt"],
        ],
    )

    doc.add_heading("6.3 Text OCR", level=2)
    doc.add_paragraph(
        "Text fields use the pre-trained PaddleOCR Arabic model at inference time. "
        "No separate fine-tuning notebook is included in the repository. "
        "Field-specific logic in field_ocr.py handles RTL stitching, date fields, and serial numbers."
    )

    doc.add_heading("7. Accuracy and Evaluation", level=1)
    doc.add_paragraph(
        "Accuracy is measured at three levels: field localization (box detection), digit detection "
        "(per-digit boxes), and end-to-end ID string match. Text OCR accuracy is validated manually "
        "through the review UI rather than a committed benchmark."
    )

    doc.add_heading("7.1 Field Localization (validation set)", level=2)
    doc.add_paragraph(
        "From training logs on the Thndr valid set (~193 images). Best late-epoch all-class metrics:"
    )
    add_table(
        doc,
        ["Metric", "Approximate value"],
        [
            ["Precision", "~0.95"],
            ["Recall", "~0.92"],
            ["mAP@0.5", "~0.95"],
            ["mAP@0.5:0.95", "~0.56"],
        ],
    )

    doc.add_heading("7.2 Digit Detection (validation set)", level=2)
    doc.add_paragraph("From cro4 valid set (849 images, early training epochs):")
    add_table(
        doc,
        ["Metric", "Approximate value"],
        [
            ["Precision", "~0.999"],
            ["Recall", "~0.999"],
            ["mAP@0.5", "~0.995"],
            ["mAP@0.5:0.95", "~0.92"],
        ],
    )

    doc.add_heading("7.3 End-to-End ID String Match", level=2)
    doc.add_paragraph(
        "The digit training notebook evaluates exact 14-digit string match on valid and test splits. "
        "This is the most important metric for the national ID number field. Reported metrics include:"
    )
    for item in [
        "Exact match rate — full 14-digit string correct",
        "14-box rate — exactly 14 digit boxes detected",
        "Character accuracy — per-digit correctness",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "Re-run train_nid_digits.ipynb with RUN_EXACT=True or consult MLflow for exact figures."
    )

    doc.add_heading("7.4 Full Pipeline (16 fields)", level=2)
    doc.add_paragraph(
        "End-to-end field accuracy depends on both localization and OCR/digit reading. "
        "Some fields (IssueDate, HusbandName, Front) are frequently missed on real-world photos; "
        "the web UI marks missing fields and allows manual correction before submission."
    )

    doc.add_heading("8. NID Decoding", level=1)
    doc.add_paragraph("A valid 14-digit Egyptian national ID encodes:")
    add_table(
        doc,
        ["Position", "Meaning"],
        [
            ["Digit 1", "Century (2 = 1900s, 3 = 2000s)"],
            ["Digits 2–7", "Birth date (YYMMDD)"],
            ["Digits 8–9", "Governorate code (e.g. 88 = Born abroad)"],
            ["Digits 10–13", "Serial number (digit 13: odd = male, even = female)"],
            ["Digit 14", "Check digit (structure validated; official MOI algorithm not verified)"],
        ],
    )

    doc.add_heading("9. Inference Configuration", level=1)
    add_table(
        doc,
        ["Setting", "Default"],
        [
            ["Localization confidence", "0.2"],
            ["Localization image size", "1280 px"],
            ["Digit confidence", "0.25"],
            ["Digit image size", "640 px"],
            ["OCR language", "Arabic (ar)"],
            ["Bundled localization weights", "yolo+ocr/weights/best_snapshot.pt"],
            ["Bundled digit weights", "yolo+ocr/weights/digit_best.pt"],
        ],
    )

    doc.add_heading("10. Limitations", level=1)
    for item in [
        "Accuracy drops with blur, glare, rotation, and partial card photos.",
        "PaddleOCR on CPU requires PaddlePaddle 3.2.2; version 3.3.x can crash with oneDNN errors.",
        "Check digit validation is structural only — the official MOI checksum is not implemented.",
        "Front/Back marker fields may show detection status text rather than OCR content.",
        "Real national ID images should not be uploaded to public demos without redaction.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("11. References", level=1)
    for item in [
        "Thndr National Card — Roboflow Universe",
        "cro4 digit dataset — Roboflow Universe",
        "Ultralytics YOLO documentation",
        "PaddleOCR / PaddlePaddle documentation",
        "Repository README and training notebooks",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    return doc


def build_user_guide() -> Document:
    doc = Document()
    set_normal_style(doc)
    add_title_page(doc, "Egyptian National ID OCR", "User Guide")

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This application helps you upload photos of an Egyptian national ID card, automatically "
        "extract 16 data fields, review and correct the results, and optionally save them to a "
        "Google Spreadsheet."
    )
    notice = doc.add_paragraph()
    notice.add_run("Privacy notice: ").bold = True
    notice.add_run(
        "Use sample or redacted card images only in demos. "
        "Do not upload real national ID cards to untrusted or public environments."
    )

    doc.add_heading("2. Before You Start", level=1)
    doc.add_paragraph("You need:")
    for item in [
        "A modern web browser (Chrome, Edge, or Firefox)",
        "Clear photos of the front (required) and back (recommended) of the card",
        "The application running locally at http://127.0.0.1:8000",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Start the server (from the project folder):")
    code = doc.add_paragraph("uvicorn web.app:app --reload --host 127.0.0.1 --port 8000")
    code.runs[0].font.name = "Consolas"
    doc.add_paragraph(
        "On first launch, wait until the status message no longer says “Loading models…” "
        "(this can take 1–2 minutes)."
    )

    doc.add_heading("3. Step-by-Step Workflow", level=1)

    doc.add_heading("Step 1 — Upload card images", level=2)
    for item in [
        "Click or drag an image onto Upload front side of the card (required).",
        "Click or drag an image onto Upload back side of the card (recommended).",
        "A preview appears after each upload.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Step 2 — Extract fields", level=2)
    for item in [
        "Click Extract fields (enabled after a front image is uploaded).",
        "Wait while the system detects fields and runs OCR (about 10–30 seconds on CPU).",
        "When complete, the status shows: “Extraction complete. Review and edit any fields below.”",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Step 3 — Review results", level=2)
    doc.add_paragraph("Saved images: annotated front/back photos with detection boxes (click to enlarge).")
    doc.add_paragraph("Form sections and fields:")
    add_table(
        doc,
        ["Section", "Fields"],
        [
            ["Identity", "First name, Last name, Husband name, Gender, Religion, Status"],
            ["ID number", "14-digit National ID"],
            ["Decoded from ID", "Birth date, governorate, gender from ID (read-only)"],
            ["Dates & serial", "Issue date, Expiry date, Serial number"],
            ["Address", "Address line 1, Address line 2"],
            ["Employment", "Job 1, Job 2"],
            ["Card sides", "Front, Back markers"],
        ],
    )
    doc.add_paragraph("Field helper messages:")
    add_table(
        doc,
        ["Message", "Meaning"],
        [
            ["Not detected — please fill in", "Field was not found; type the value manually"],
            ["From front/back (conf X%)", "Value extracted automatically with detection confidence"],
        ],
    )

    doc.add_heading("Step 4 — Edit missing or wrong values", level=2)
    for item in [
        "Click any text field and type corrections.",
        "Arabic fields support right-to-left text.",
        "The ID field accepts Western digits (0–9).",
        "The Decoded from ID panel updates when the ID number changes.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Step 5 — Submit to spreadsheet (optional)", level=2)
    for item in [
        "After a successful extract, Submit to spreadsheet becomes active.",
        "Click it to send the current form values (not images) to Google Sheets.",
        "Success message: “Submitted to spreadsheet.”",
        "Each submit adds a new row; it does not update previous rows.",
        "If submit fails, an administrator must configure GOOGLE_SHEETS_WEBHOOK_URL on the server.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("4. Tips for Best Results", level=1)
    add_table(
        doc,
        ["Tip", "Why"],
        [
            ["Use good lighting, minimal glare", "Improves detection and OCR accuracy"],
            ["Fill the frame with the card", "Helps find all field regions"],
            ["Upload both front and back", "Back holds Gender, Religion, Expiry, Job fields"],
            ["Keep the card flat and in focus", "Reduces Arabic OCR errors"],
            ["Always review before submitting", "OCR is not 100% accurate"],
        ],
    )

    doc.add_heading("5. Troubleshooting", level=1)
    add_table(
        doc,
        ["Problem", "Solution"],
        [
            ["Extract fields is disabled", "Upload a front image first"],
            ["Loading models… never finishes", "Refresh the page; ensure the server is running"],
            ["Extractor not ready", "Wait for models to finish loading"],
            ["Many empty fields", "Use a clearer photo; upload the back; fill fields manually"],
            ["Submit fails", "Ask admin to set GOOGLE_SHEETS_WEBHOOK_URL"],
            ["Wrong decoded ID fields", "Correct the 14-digit ID first; decode panel updates automatically"],
        ],
    )

    doc.add_heading("6. What Gets Saved Locally", level=1)
    doc.add_paragraph("Each extraction saves files under web/uploads/<run_id>/:")
    for item in [
        "front.jpg / back.jpg — original uploads",
        "front_annotated.jpg / back_annotated.jpg — images with detection boxes",
        "crops/front/ and crops/back/ — individual field crops",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7. Field Glossary", level=1)
    add_table(
        doc,
        ["Field", "Description", "Typical side"],
        [
            ["First_Name", "Given name", "Front"],
            ["Last_Name", "Full family name", "Front"],
            ["HusbandName", "Husband's name (if applicable)", "Front"],
            ["ID", "14-digit national ID number", "Front"],
            ["Serial_Num", "Card serial number", "Front"],
            ["Add1 / Add2", "Address lines", "Front"],
            ["Gender", "Male / Female", "Back"],
            ["Religion", "Religion", "Back"],
            ["Status", "Marital status", "Back"],
            ["IssueDate / ExpDate", "Issue and expiry dates", "Back"],
            ["Job1 / Job2", "Occupation", "Back"],
        ],
    )

    doc.add_heading("8. Support", level=1)
    doc.add_paragraph(
        "For technical or model questions, refer to the Technical Documentation. "
        "For server or spreadsheet setup, contact your system administrator."
    )

    return doc


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    technical_path = DOCS_DIR / "Egyptian_NID_OCR_Technical_Documentation.docx"
    user_path = DOCS_DIR / "Egyptian_NID_OCR_User_Guide.docx"
    build_technical_doc().save(technical_path)
    build_user_guide().save(user_path)
    print(f"Wrote {technical_path}")
    print(f"Wrote {user_path}")


if __name__ == "__main__":
    main()
